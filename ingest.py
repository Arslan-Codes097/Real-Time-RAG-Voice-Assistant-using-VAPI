import os
import sys
import uuid
import glob
from typing import List
from dotenv import load_dotenv

load_dotenv()

if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from fastembed import TextEmbedding
from pypdf import PdfReader
import docx

DOCS_DIRECTORY = os.path.join(os.path.dirname(__file__), "sample_docs")

SUPABASE_URL = os.getenv("SUPABASE_URL", "")
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "")

from langchain_huggingface import HuggingFaceEndpointEmbeddings
from langchain_community.vectorstores import SupabaseVectorStore
from supabase import create_client

def get_embeddings():
    hf_token = os.getenv("HF_TOKEN", "")
    if not hf_token:
        print("HF_TOKEN is missing! Please add it to your .env file.")
        return None
    return HuggingFaceEndpointEmbeddings(
        model="BAAI/bge-small-en-v1.5",
        huggingfacehub_api_token=hf_token
    )


import pdfplumber

def load_pdf(file_path: str) -> List[Document]:
    docs = []
    filename = os.path.basename(file_path)
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # Using layout=True to preserve spatial arrangement and tables
                text = page.extract_text(layout=True) or page.extract_text() or ""
                if text.strip():
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={"source": filename, "source_file": filename, "page": i + 1}
                        )
                    )
        return docs
    except Exception as e:
        print(f"pdfplumber failed for {filename} ({e}), falling back to pypdf...")
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": filename, "source_file": filename, "page": i + 1}
                    )
                )
        return docs

def load_docx(file_path: str) -> List[Document]:
    doc = docx.Document(file_path)
    full_text = []
    filename = os.path.basename(file_path)
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip():
                    row_data.append(cell.text.strip())
            if row_data:
                full_text.append(" | ".join(row_data))
                
    content = "\n".join(full_text)
    if content.strip():
        return [Document(page_content=content, metadata={"source": filename, "source_file": filename})]
    return []

def load_txt(file_path: str) -> List[Document]:
    filename = os.path.basename(file_path)
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    if content.strip():
        return [Document(page_content=content, metadata={"source": filename, "source_file": filename})]
    return []

def load_all_documents(folder_path: str) -> List[Document]:
    documents = []
    for file_path in glob.glob(os.path.join(folder_path, "*")):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            print(f"📄 Loading PDF: {os.path.basename(file_path)}")
            documents.extend(load_pdf(file_path))
        elif ext == ".docx":
            print(f"📝 Loading DOCX: {os.path.basename(file_path)}")
            documents.extend(load_docx(file_path))
        elif ext in [".txt", ".md"]:
            print(f"📑 Loading Text: {os.path.basename(file_path)}")
            documents.extend(load_txt(file_path))
    return documents

def ingest_documents():
    print("🚀 Starting Standard LangChain Document Ingestion...")
    raw_docs = load_all_documents(DOCS_DIRECTORY)
    print(f"✅ Loaded {len(raw_docs)} document sections/pages.")

    if not raw_docs:
        print("⚠️ No valid documents found in sample_docs directory.")
        return

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=750,
        chunk_overlap=100,
        separators=["\n\n", "\n", " ", ""]
    )
    chunks = text_splitter.split_documents(raw_docs)
    print(f"🧩 Split into {len(chunks)} text chunks.")
    
    # Add index to metadata
    for i, chunk in enumerate(chunks):
        if "source_file" not in chunk.metadata:
            chunk.metadata["source_file"] = chunk.metadata.get("source", "Document")
        chunk.metadata["chunk_index"] = i

    if SUPABASE_URL and SUPABASE_KEY and "your_supabase" not in SUPABASE_URL:
        print("⚡ Upserting vector chunks into Supabase Cloud Vector DB using LangChain...")
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
        embeddings = get_embeddings()
        
        # We can use SupabaseVectorStore.from_documents
        vector_store = SupabaseVectorStore(
            client=supabase,
            embedding=embeddings,
            table_name="documents",
            query_name="match_documents",
        )
        
        # Batch insert
        batch_size = 50
        import uuid
        for idx in range(0, len(chunks), batch_size):
            batch = chunks[idx:idx + batch_size]
            batch_ids = [str(uuid.uuid4()) for _ in batch]
            vector_store.add_documents(batch, ids=batch_ids)
            print(f"  -> Uploaded batch {idx//batch_size + 1}/{(len(chunks)-1)//batch_size + 1}")
            
        print("🎉 Ingestion into Supabase Vector DB Complete!")

if __name__ == "__main__":
    ingest_documents()
